import ast
import re
from dataclasses import dataclass
from pathlib import Path
import zipfile

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings
from app.utils.file_utils import safe_relative_path, should_ignore_dir, should_ignore_file


@dataclass
class ParsedFile:
    file_path: str
    content: str


@dataclass
class DocumentChunk:
    content: str
    metadata: dict


LANGUAGE_BY_EXTENSION = {
    ".py": "python",
    ".js": "javascript",
    ".jsx": "javascript",
    ".mjs": "javascript",
    ".cjs": "javascript",
    ".ts": "typescript",
    ".tsx": "typescript",
    ".java": "java",
    ".go": "go",
    ".rs": "rust",
    ".c": "c",
    ".cc": "cpp",
    ".cpp": "cpp",
    ".cxx": "cpp",
    ".h": "c",
    ".hpp": "cpp",
    ".cs": "csharp",
    ".php": "php",
    ".rb": "ruby",
    ".kt": "kotlin",
    ".kts": "kotlin",
    ".swift": "swift",
    ".vue": "vue",
    ".svelte": "svelte",
    ".html": "html",
    ".htm": "html",
    ".css": "css",
    ".scss": "scss",
    ".sass": "sass",
    ".less": "less",
    ".sh": "shell",
    ".bash": "shell",
    ".zsh": "shell",
    ".ps1": "powershell",
    ".bat": "batch",
    ".cmd": "batch",
    ".md": "markdown",
    ".mdx": "markdown",
    ".rst": "text",
    ".adoc": "text",
    ".json": "json",
    ".jsonl": "json",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".toml": "toml",
    ".ini": "ini",
    ".cfg": "config",
    ".conf": "config",
    ".properties": "properties",
    ".xml": "xml",
    ".sql": "sql",
    ".graphql": "graphql",
    ".gql": "graphql",
    ".proto": "protobuf",
    ".prisma": "prisma",
    ".gradle": "gradle",
    ".csv": "csv",
    ".tsv": "tsv",
    ".txt": "text",
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
}

LANGUAGE_BY_FILENAME = {
    "dockerfile": "dockerfile",
    "containerfile": "dockerfile",
    "makefile": "makefile",
    "cmakelists.txt": "cmake",
    "requirements.txt": "requirements",
    "pipfile": "toml",
    "gemfile": "ruby",
    "rakefile": "ruby",
    "cargo.toml": "toml",
    "go.mod": "go",
    "go.sum": "text",
    "package.json": "json",
    "pyproject.toml": "toml",
    "setup.py": "python",
}

TEXT_DECODINGS = ("utf-8", "utf-8-sig", "gb18030", "cp936", "cp1252", "latin-1")
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
SAFE_DOCUMENT_CONTAINER_COMPRESSION = {zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED}


def looks_binary(data: bytes) -> bool:
    if b"\x00" in data:
        return True
    if not data:
        return False
    sample = data[:4096]
    control_bytes = sum(1 for byte in sample if byte < 9 or (13 < byte < 32))
    return control_bytes / len(sample) > 0.18


def read_text_content(path: Path) -> str | None:
    try:
        data = path.read_bytes()
    except OSError:
        return None

    if looks_binary(data):
        return None

    for encoding in TEXT_DECODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return None


def office_container_is_within_limits(path: Path) -> bool:
    settings = get_settings()
    try:
        with zipfile.ZipFile(path, mode="r") as archive:
            entries = archive.infolist()
            if len(entries) > settings.max_document_container_files:
                return False
            total_uncompressed = 0
            total_compressed = 0
            for entry in entries:
                if (
                    entry.flag_bits & (0x1 | 0x20 | 0x40 | 0x2000)
                    or entry.compress_type not in SAFE_DOCUMENT_CONTAINER_COMPRESSION
                    or entry.file_size < 0
                    or entry.compress_size < 0
                ):
                    return False
                if entry.is_dir():
                    continue
                total_uncompressed += entry.file_size
                total_compressed += entry.compress_size
                if total_uncompressed > settings.max_document_uncompressed_bytes:
                    return False
                if entry.file_size:
                    if entry.compress_size <= 0:
                        return False
                    if (
                        entry.file_size / entry.compress_size
                        > settings.max_document_compression_ratio
                    ):
                        return False
            if total_uncompressed:
                if total_compressed <= 0:
                    return False
                if (
                    total_uncompressed / total_compressed
                    > settings.max_document_compression_ratio
                ):
                    return False
    except (OSError, RuntimeError, zipfile.BadZipFile, zipfile.LargeZipFile):
        return False
    return True


def read_pdf_content(path: Path) -> str | None:
    try:
        from pypdf import PdfReader
    except ImportError:
        return None

    try:
        reader = PdfReader(str(path))
        settings = get_settings()
        if len(reader.pages) > settings.max_document_pages:
            return None
        pages: list[str] = []
        extracted_characters = 0
        for page in reader.pages:
            content = page.extract_text() or ""
            extracted_characters += len(content)
            if extracted_characters > settings.max_document_extracted_characters:
                return None
            pages.append(content)
    except Exception:
        return None
    return "\n\n".join(page.strip() for page in pages if page.strip()) or None


def read_docx_content(path: Path) -> str | None:
    try:
        from docx import Document
    except ImportError:
        return None

    if not office_container_is_within_limits(path):
        return None

    try:
        settings = get_settings()
        document = Document(str(path))
        extracted_characters = 0
        paragraphs: list[str] = []
        for paragraph in document.paragraphs:
            content = paragraph.text.strip()
            if not content:
                continue
            extracted_characters += len(content)
            if extracted_characters > settings.max_document_extracted_characters:
                return None
            paragraphs.append(content)
        table_rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    content = " | ".join(cells)
                    extracted_characters += len(content)
                    if extracted_characters > settings.max_document_extracted_characters:
                        return None
                    table_rows.append(content)
    except Exception:
        return None
    return "\n".join(paragraphs + table_rows).strip() or None


def read_xlsx_content(path: Path) -> str | None:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return None

    if not office_container_is_within_limits(path):
        return None

    workbook = None
    try:
        settings = get_settings()
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        if len(workbook.worksheets) > settings.max_document_sheets:
            return None
        sections: list[str] = []
        total_rows = 0
        total_cells = 0
        extracted_characters = 0
        for sheet in workbook.worksheets:
            rows: list[str] = [f"# Sheet: {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                total_rows += 1
                total_cells += len(row)
                if (
                    total_rows > settings.max_document_rows
                    or total_cells > settings.max_document_cells
                ):
                    return None
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    content = " | ".join(values)
                    extracted_characters += len(content)
                    if extracted_characters > settings.max_document_extracted_characters:
                        return None
                    rows.append(content)
            if len(rows) > 1:
                sections.append("\n".join(rows))
    except Exception:
        return None
    finally:
        if workbook is not None:
            workbook.close()
    return "\n\n".join(sections).strip() or None


def read_supported_file(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf_content(path)
    if suffix == ".docx":
        return read_docx_content(path)
    if suffix == ".xlsx":
        return read_xlsx_content(path)
    return read_text_content(path)


def read_repository_files(repo_path: Path) -> list[ParsedFile]:
    settings = get_settings()
    parsed_files: list[ParsedFile] = []
    total_bytes = 0

    for path in repo_path.rglob("*"):
        if any(should_ignore_dir(parent) for parent in path.relative_to(repo_path).parents):
            continue
        if path.is_symlink():
            continue
        if path.is_dir():
            continue
        if should_ignore_file(path):
            continue
        try:
            size = path.stat().st_size
        except OSError:
            continue
        if size > settings.max_file_size_bytes:
            continue
        if len(parsed_files) >= settings.max_repository_files:
            break
        if total_bytes + size > settings.max_repository_bytes:
            continue

        content = read_supported_file(path)
        if content is None:
            continue

        if content.strip():
            parsed_files.append(
                ParsedFile(file_path=safe_relative_path(path, repo_path), content=content)
            )
            total_bytes += size

    return parsed_files


def get_language(file_path: str) -> str:
    path = Path(file_path)
    return LANGUAGE_BY_FILENAME.get(path.name.lower(), LANGUAGE_BY_EXTENSION.get(path.suffix.lower(), "text"))


def get_line_range(lines: list[str], start_line: int, end_line: int) -> str:
    start_index = max(start_line - 1, 0)
    end_index = min(end_line, len(lines))
    return "\n".join(lines[start_index:end_index]).strip()


def make_chunk(
    *,
    content: str,
    repository_id: str,
    file_path: str,
    chunk_index: int,
    language: str,
    start_line: int,
    end_line: int,
    symbol_name: str = "",
    symbol_type: str = "text",
    parent_symbol: str = "",
) -> DocumentChunk:
    return DocumentChunk(
        content=content,
        metadata={
            "repository_id": repository_id,
            "file_path": file_path,
            "chunk_index": chunk_index,
            "start_line": start_line,
            "end_line": end_line,
            "language": language,
            "symbol_name": symbol_name,
            "symbol_type": symbol_type,
            "parent_symbol": parent_symbol,
        },
    )


def split_large_block(
    *,
    content: str,
    repository_id: str,
    file_path: str,
    language: str,
    chunk_index: int,
    start_line: int,
    symbol_name: str,
    symbol_type: str,
    parent_symbol: str = "",
) -> list[DocumentChunk]:
    settings = get_settings()
    block_lines = content.splitlines()
    if len(content) <= settings.chunk_size * 2:
        return [
            make_chunk(
                content=content,
                repository_id=repository_id,
                file_path=file_path,
                chunk_index=chunk_index,
                language=language,
                start_line=start_line,
                end_line=start_line + max(len(block_lines), 1) - 1,
                symbol_name=symbol_name,
                symbol_type=symbol_type,
                parent_symbol=parent_symbol,
            )
        ]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks: list[DocumentChunk] = []
    cursor_line = start_line
    for offset, text in enumerate(splitter.split_text(content)):
        line_count = max(text.count("\n") + 1, 1)
        chunks.append(
            make_chunk(
                content=text,
                repository_id=repository_id,
                file_path=file_path,
                chunk_index=chunk_index + offset,
                language=language,
                start_line=cursor_line,
                end_line=cursor_line + line_count - 1,
                symbol_name=symbol_name,
                symbol_type=symbol_type,
                parent_symbol=parent_symbol,
            )
        )
        cursor_line += max(line_count - 1, 1)
    return chunks


def split_python_file(parsed_file: ParsedFile, repository_id: str) -> list[DocumentChunk]:
    lines = parsed_file.content.splitlines()
    chunks: list[DocumentChunk] = []
    seen_ranges: set[tuple[int, int]] = set()

    try:
        tree = ast.parse(parsed_file.content)
    except SyntaxError:
        return []

    class_parent: dict[ast.AST, str] = {}
    for node in ast.walk(tree):
        if isinstance(node, ast.ClassDef):
            for child in node.body:
                class_parent[child] = node.name

    interesting_nodes = [
        node
        for node in ast.walk(tree)
        if isinstance(node, (ast.ClassDef, ast.FunctionDef, ast.AsyncFunctionDef))
        and hasattr(node, "lineno")
        and hasattr(node, "end_lineno")
    ]
    interesting_nodes.sort(key=lambda node: (node.lineno, getattr(node, "end_lineno", node.lineno)))

    for node in interesting_nodes:
        start_line = int(node.lineno)
        end_line = int(node.end_lineno or node.lineno)
        if (start_line, end_line) in seen_ranges:
            continue
        seen_ranges.add((start_line, end_line))

        content = get_line_range(lines, start_line, end_line)
        if not content:
            continue
        parent_symbol = class_parent.get(node, "")
        if isinstance(node, ast.ClassDef):
            symbol_type = "class"
        elif parent_symbol:
            symbol_type = "method"
        else:
            symbol_type = "function"
        chunks.extend(
            split_large_block(
                content=content,
                repository_id=repository_id,
                file_path=parsed_file.file_path,
                language="python",
                chunk_index=len(chunks),
                start_line=start_line,
                symbol_name=getattr(node, "name", ""),
                symbol_type=symbol_type,
                parent_symbol=parent_symbol,
            )
        )

    return chunks


CODE_BLOCK_PATTERNS = {
    "javascript": re.compile(
        r"^\s*(export\s+)?(async\s+)?(function\s+([A-Za-z_$][\w$]*)|class\s+([A-Za-z_$][\w$]*)|const\s+([A-Za-z_$][\w$]*)\s*=\s*(async\s*)?\([^)]*\)\s*=>)",
        re.MULTILINE,
    ),
    "typescript": re.compile(
        r"^\s*(export\s+)?(async\s+)?(function\s+([A-Za-z_$][\w$]*)|class\s+([A-Za-z_$][\w$]*)|interface\s+([A-Za-z_$][\w$]*)|type\s+([A-Za-z_$][\w$]*)|const\s+([A-Za-z_$][\w$]*)\s*=\s*(async\s*)?\([^)]*\)\s*=>)",
        re.MULTILINE,
    ),
    "java": re.compile(
        r"^\s*((public|private|protected|static|final|abstract)\s+)*(class|interface|enum)\s+([A-Za-z_][\w]*)|^\s*((public|private|protected|static|final|synchronized)\s+)+[\w<>\[\], ?]+\s+([A-Za-z_][\w]*)\s*\(",
        re.MULTILINE,
    ),
    "go": re.compile(
        r"^\s*(func\s+(\([^)]+\)\s*)?([A-Za-z_][\w]*)\s*\(|type\s+([A-Za-z_][\w]*)\s+(struct|interface))",
        re.MULTILINE,
    ),
    "rust": re.compile(
        r"^\s*(pub\s+)?(async\s+)?(fn\s+([A-Za-z_][\w]*)|struct\s+([A-Za-z_][\w]*)|enum\s+([A-Za-z_][\w]*)|trait\s+([A-Za-z_][\w]*)|impl\b)",
        re.MULTILINE,
    ),
    "c": re.compile(
        r"^\s*((static|inline|extern)\s+)*[\w\*\s]+\s+([A-Za-z_][\w]*)\s*\([^;]*\)\s*\{|^\s*(typedef\s+)?struct\s+([A-Za-z_][\w]*)?",
        re.MULTILINE,
    ),
    "cpp": re.compile(
        r"^\s*((template\s*<[^>]+>\s*)?((class|struct|enum)\s+([A-Za-z_][\w]*)|[\w:<>\*&\s]+\s+([A-Za-z_~][\w~]*)\s*\([^;]*\)\s*(const\s*)?\{))",
        re.MULTILINE,
    ),
    "csharp": re.compile(
        r"^\s*((public|private|protected|internal|static|sealed|abstract|async)\s+)*(class|interface|enum|struct|record)\s+([A-Za-z_][\w]*)|^\s*((public|private|protected|internal|static|async|virtual|override)\s+)+[\w<>\[\], ?]+\s+([A-Za-z_][\w]*)\s*\(",
        re.MULTILINE,
    ),
    "php": re.compile(
        r"^\s*((public|private|protected|static|final|abstract)\s+)*(function\s+([A-Za-z_][\w]*)|class\s+([A-Za-z_][\w]*)|interface\s+([A-Za-z_][\w]*)|trait\s+([A-Za-z_][\w]*))",
        re.MULTILINE,
    ),
    "ruby": re.compile(
        r"^\s*(class|module|def)\s+([A-Za-z_][\w:!?=]*)",
        re.MULTILINE,
    ),
    "kotlin": re.compile(
        r"^\s*((public|private|protected|internal|open|data|sealed|abstract|suspend)\s+)*(class|interface|object|fun)\s+([A-Za-z_][\w]*)",
        re.MULTILINE,
    ),
    "swift": re.compile(
        r"^\s*((public|private|fileprivate|internal|open|static|final)\s+)*(class|struct|enum|protocol|func)\s+([A-Za-z_][\w]*)",
        re.MULTILINE,
    ),
}


def find_block_end(lines: list[str], start_line: int) -> int:
    brace_balance = 0
    saw_brace = False
    for index in range(start_line - 1, len(lines)):
        line = lines[index]
        brace_balance += line.count("{") - line.count("}")
        if "{" in line:
            saw_brace = True
        if saw_brace and brace_balance <= 0:
            return index + 1
        if not saw_brace and index > start_line and not line.strip():
            return index
    return len(lines)


def extract_symbol_name(match: re.Match[str]) -> str:
    for group in reversed(match.groups()):
        if group and re.match(r"^[A-Za-z_$][\w$]*$", group):
            return group
    return ""


def split_brace_language_file(parsed_file: ParsedFile, repository_id: str, language: str) -> list[DocumentChunk]:
    pattern = CODE_BLOCK_PATTERNS.get(language)
    if not pattern:
        return []
    lines = parsed_file.content.splitlines()
    chunks: list[DocumentChunk] = []
    seen_starts: set[int] = set()

    for match in pattern.finditer(parsed_file.content):
        start_line = parsed_file.content[: match.start()].count("\n") + 1
        if start_line in seen_starts:
            continue
        seen_starts.add(start_line)
        end_line = find_block_end(lines, start_line)
        content = get_line_range(lines, start_line, end_line)
        if not content:
            continue
        first_line = lines[start_line - 1].lower() if start_line <= len(lines) else ""
        if any(word in first_line for word in ("class ", "interface ", "enum ", "struct")):
            symbol_type = "class"
        elif language == "java":
            symbol_type = "method"
        elif language in {"javascript", "typescript"} and "export" in first_line:
            symbol_type = "export"
        else:
            symbol_type = "function"
        chunks.extend(
            split_large_block(
                content=content,
                repository_id=repository_id,
                file_path=parsed_file.file_path,
                language=language,
                chunk_index=len(chunks),
                start_line=start_line,
                symbol_name=extract_symbol_name(match),
                symbol_type=symbol_type,
            )
        )

    return chunks


def split_markdown_file(parsed_file: ParsedFile, repository_id: str) -> list[DocumentChunk]:
    lines = parsed_file.content.splitlines()
    heading_lines = [index + 1 for index, line in enumerate(lines) if line.startswith("#")]
    if not heading_lines:
        return []

    chunks: list[DocumentChunk] = []
    for offset, start_line in enumerate(heading_lines):
        end_line = heading_lines[offset + 1] - 1 if offset + 1 < len(heading_lines) else len(lines)
        content = get_line_range(lines, start_line, end_line)
        if not content:
            continue
        heading = lines[start_line - 1].lstrip("#").strip()
        chunks.extend(
            split_large_block(
                content=content,
                repository_id=repository_id,
                file_path=parsed_file.file_path,
                language="markdown",
                chunk_index=len(chunks),
                start_line=start_line,
                symbol_name=heading,
                symbol_type="section",
            )
        )
    return chunks


def split_fallback_file(parsed_file: ParsedFile, repository_id: str, language: str) -> list[DocumentChunk]:
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks: list[DocumentChunk] = []
    cursor_line = 1
    for text in splitter.split_text(parsed_file.content):
        line_count = max(text.count("\n") + 1, 1)
        chunks.append(
            make_chunk(
                content=text,
                repository_id=repository_id,
                file_path=parsed_file.file_path,
                chunk_index=len(chunks),
                language=language,
                start_line=cursor_line,
                end_line=cursor_line + line_count - 1,
                symbol_type="text",
            )
        )
        cursor_line += max(line_count - 1, 1)
    return chunks


def split_files_into_chunks(files: list[ParsedFile], repository_id: str) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    for parsed_file in files:
        language = get_language(parsed_file.file_path)
        file_chunks: list[DocumentChunk] = []
        if language == "python":
            file_chunks = split_python_file(parsed_file, repository_id)
        elif language in CODE_BLOCK_PATTERNS:
            file_chunks = split_brace_language_file(parsed_file, repository_id, language)
        elif language == "markdown":
            file_chunks = split_markdown_file(parsed_file, repository_id)

        if not file_chunks:
            file_chunks = split_fallback_file(parsed_file, repository_id, language)

        for chunk in file_chunks:
            chunk.metadata["chunk_index"] = len([item for item in chunks if item.metadata["file_path"] == parsed_file.file_path])
            chunks.append(chunk)

    return chunks
