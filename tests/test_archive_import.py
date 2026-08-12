from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
import json
from pathlib import Path
from threading import Event
import stat
from urllib.parse import quote
import zipfile

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
import pytest
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

import app.main as main_module
from app.config import get_settings
from app.main import app
from app.security import auth
from app.services.repository_catalog import RepositoryCatalogService
from app.services.vector_store import IndexResult


def configure_archive_test(monkeypatch, tmp_path):
    settings = get_settings()
    monkeypatch.setattr(settings, "repos_dir", tmp_path / "repos")
    monkeypatch.setattr(settings, "chroma_dir", tmp_path / "chroma")
    monkeypatch.setattr(settings, "app_api_key", "")
    monkeypatch.setattr(settings, "admin_username", "")
    monkeypatch.setattr(settings, "admin_password_hash", "")
    monkeypatch.setattr(settings, "auth_session_secret", "")
    monkeypatch.setattr(settings, "rate_limit_max_requests", 0)
    monkeypatch.setattr(settings, "security_audit_enabled", True)
    monkeypatch.setattr(settings, "security_audit_log_path", tmp_path / "security-audit.jsonl")
    monkeypatch.setattr(settings, "max_file_size_bytes", 1_000_000)
    monkeypatch.setattr(settings, "max_repository_files", 100)
    monkeypatch.setattr(settings, "max_repository_bytes", 5_000_000)
    monkeypatch.setattr(settings, "max_archive_upload_bytes", 2_000_000)
    monkeypatch.setattr(settings, "max_archive_files", 100)
    monkeypatch.setattr(settings, "max_archive_directories", 100)
    monkeypatch.setattr(settings, "max_archive_uncompressed_bytes", 5_000_000)
    monkeypatch.setattr(settings, "max_archive_compression_ratio", 100.0)
    monkeypatch.setattr(settings, "max_archive_path_depth", 12)
    monkeypatch.setattr(settings, "max_archive_path_length", 256)
    monkeypatch.setattr(settings, "max_archive_name_length", 255)
    monkeypatch.setattr(settings, "repository_import_timeout_seconds", 30)
    settings.repos_dir.mkdir(parents=True, exist_ok=True)
    settings.chroma_dir.mkdir(parents=True, exist_ok=True)
    auth.clear_login_attempts()
    return settings


def fake_index_result(_repository_id, chunks):
    return IndexResult(
        chunks_indexed=len(chunks),
        chunks_written=len(chunks),
        files_indexed=len({chunk.metadata["file_path"] for chunk in chunks}),
        index_cached=False,
        changed_files_count=len({chunk.metadata["file_path"] for chunk in chunks}),
        removed_files_count=0,
    )


def make_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Upload guide", level=1)
    document.add_paragraph("The ZIP importer keeps this Word document searchable.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Roadmap"
    sheet.append(["priority", "feature"])
    sheet.append(["P0", "safe ZIP import"])
    buffer = BytesIO()
    workbook.save(buffer)
    workbook.close()
    return buffer.getvalue()


def make_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 40 250 Td (Searchable archive PDF) Tj ET")
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_zip(
    entries: dict[str, bytes | str],
    *,
    compression: int = zipfile.ZIP_DEFLATED,
) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=compression) as archive:
        for name, content in entries.items():
            data = content.encode("utf-8") if isinstance(content, str) else content
            archive.writestr(name, data)
    return buffer.getvalue()


def make_symlink_zip() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w") as archive:
        entry = zipfile.ZipInfo("linked.py")
        entry.create_system = 3
        entry.external_attr = (stat.S_IFLNK | 0o777) << 16
        archive.writestr(entry, "../outside.py")
    return buffer.getvalue()


def mark_first_entry_encrypted(data: bytes) -> bytes:
    payload = bytearray(data)
    local_header = payload.find(b"PK\x03\x04")
    central_header = payload.find(b"PK\x01\x02")
    assert local_header >= 0 and central_header >= 0
    local_flags = int.from_bytes(payload[local_header + 6 : local_header + 8], "little")
    central_flags = int.from_bytes(
        payload[central_header + 8 : central_header + 10], "little"
    )
    payload[local_header + 6 : local_header + 8] = (local_flags | 0x1).to_bytes(
        2, "little"
    )
    payload[central_header + 8 : central_header + 10] = (
        central_flags | 0x1
    ).to_bytes(2, "little")
    return bytes(payload)


def upload_zip(
    client: TestClient,
    payload: bytes,
    *,
    name: str = "demo.zip",
    headers: dict[str, str] | None = None,
):
    request_headers = {
        "Content-Type": "application/zip",
        "X-Archive-Name": quote(name, safe=""),
    }
    request_headers.update(headers or {})
    return client.post(
        "/repository/upload-zip",
        content=payload,
        headers=request_headers,
    )


def test_zip_upload_indexes_code_and_documents_and_saves_safe_snapshot(
    tmp_path,
    monkeypatch,
):
    settings = configure_archive_test(monkeypatch, tmp_path)
    monkeypatch.setattr(main_module, "index_chunks_incremental", fake_index_result)
    nested_zip = make_zip({"ignored.py": "print('nested')"})
    docx_bytes = make_docx_bytes()
    pdf_bytes = make_pdf_bytes()
    xlsx_bytes = make_xlsx_bytes()
    payload = make_zip(
        {
            "README.md": "# Mixed knowledge base\n",
            "src/main.py": "def health():\n    return 'ok'\n",
            "docs/guide.docx": docx_bytes,
            "docs/reference.pdf": pdf_bytes,
            "sheets/roadmap.xlsx": xlsx_bytes,
            ".env": "DATABASE_URL=postgres://secret",
            "config/secrets.json": '{"token":"must-not-index"}',
            "nested.zip": nested_zip,
        }
    )

    with TestClient(app) as client:
        response = upload_zip(client, payload, name="课程资料.zip")

    assert response.status_code == 200
    body = response.json()
    assert body["files_indexed"] == 5
    assert body["chunks_indexed"] >= 5
    repository_id = body["repository_id"]
    assert repository_id.startswith("archive-")

    repository_dir = settings.repos_dir / repository_id
    snapshot = repository_dir / "source_snapshot"
    assert (snapshot / "README.md").is_file()
    assert (snapshot / "src" / "main.py").is_file()
    assert (snapshot / "docs" / "guide.docx").read_bytes() == docx_bytes
    assert (snapshot / "docs" / "reference.pdf").read_bytes() == pdf_bytes
    assert (snapshot / "sheets" / "roadmap.xlsx").read_bytes() == xlsx_bytes
    assert not (snapshot / ".env").exists()
    assert not (snapshot / "config" / "secrets.json").exists()
    assert not (snapshot / "nested.zip").exists()

    manifest = json.loads(
        (repository_dir / ".codebase_agent" / "source_manifest.json").read_text(
            encoding="utf-8"
        )
    )
    assert manifest["source"] == "zip_upload"
    assert manifest["source_type"] == "zip_upload"
    assert manifest["source_name"] == "课程资料.zip"
    assert manifest["display_name"] == "课程资料.zip"
    assert manifest["upload_name"] == "课程资料.zip"
    assert manifest["files_indexed"] == 5
    assert set(manifest["file_paths"]) == {
        "README.md",
        "src/main.py",
        "docs/guide.docx",
        "docs/reference.pdf",
        "sheets/roadmap.xlsx",
    }

    summary = RepositoryCatalogService(settings.repos_dir).get_repository(repository_id)
    assert summary.source == "zip_upload"
    assert summary.source_name == "课程资料.zip"
    assert summary.source_type == "zip_upload"
    assert summary.display_name == "课程资料.zip"
    assert summary.upload_name == "课程资料.zip"
    assert summary.github_url is None
    assert summary.files_indexed == 5

    audit_events = [
        json.loads(line)
        for line in settings.security_audit_log_path.read_text(
            encoding="utf-8"
        ).splitlines()
    ]
    assert [(event["event"], event["outcome"]) for event in audit_events] == [
        ("repository_import", "success")
    ]
    serialized_audit = settings.security_audit_log_path.read_text(encoding="utf-8")
    assert "课程资料.zip" not in serialized_audit
    assert "must-not-index" not in serialized_audit


@pytest.mark.parametrize(
    "entry_name",
    [
        "../escape.py",
        "safe/../../escape.py",
        "/absolute.py",
        "C:/Windows/system.py",
        "\\\\server\\share\\remote.py",
        "safe\\..\\escape.py",
        "file.py:alternate-stream",
    ],
)
def test_zip_upload_rejects_unsafe_paths_without_disclosure(
    entry_name,
    tmp_path,
    monkeypatch,
):
    settings = configure_archive_test(monkeypatch, tmp_path)
    payload = make_zip({entry_name: "print('private-body-marker')"})

    with TestClient(app) as client:
        response = upload_zip(client, payload)

    assert response.status_code == 400
    assert response.json() == {
        "detail": "ZIP archive contains an unsafe entry path"
    }
    assert entry_name not in response.text
    assert "private-body-marker" not in response.text
    assert not (tmp_path / "escape.py").exists()
    assert not any(path.name.startswith(".zip-upload-") for path in settings.repos_dir.iterdir())


@pytest.mark.parametrize(
    ("payload_factory", "expected_detail"),
    [
        (
            make_symlink_zip,
            "ZIP archive contains a symbolic link or special file",
        ),
        (
            lambda: mark_first_entry_encrypted(
                make_zip({"README.md": "# encrypted marker"})
            ),
            "encrypted ZIP archives are not supported",
        ),
    ],
)
def test_zip_upload_rejects_encrypted_and_special_entries(
    payload_factory,
    expected_detail,
    tmp_path,
    monkeypatch,
):
    configure_archive_test(monkeypatch, tmp_path)

    with TestClient(app) as client:
        response = upload_zip(client, payload_factory())

    assert response.status_code == 400
    assert response.json() == {"detail": expected_detail}


@pytest.mark.parametrize(
    ("setting_name", "setting_value", "entries", "expected_detail"),
    [
        (
            "max_archive_files",
            1,
            {"one.py": "1", "two.py": "2"},
            "ZIP archive exceeds the configured file limit",
        ),
        (
            "max_archive_directories",
            1,
            {"one/two/main.py": "print('ok')"},
            "ZIP archive exceeds the configured directory limit",
        ),
        (
            "max_archive_path_depth",
            2,
            {"one/two/main.py": "print('ok')"},
            "ZIP archive exceeds the configured path-depth limit",
        ),
        (
            "max_archive_uncompressed_bytes",
            32,
            {"large.txt": "x" * 33},
            "ZIP archive exceeds the configured expanded-size limit",
        ),
    ],
)
def test_zip_upload_enforces_archive_limits(
    setting_name,
    setting_value,
    entries,
    expected_detail,
    tmp_path,
    monkeypatch,
):
    settings = configure_archive_test(monkeypatch, tmp_path)
    monkeypatch.setattr(settings, setting_name, setting_value)
    payload = make_zip(entries, compression=zipfile.ZIP_STORED)

    with TestClient(app) as client:
        response = upload_zip(client, payload)

    assert response.status_code == 413
    assert response.json() == {"detail": expected_detail}


def test_zip_upload_rejects_high_compression_ratio(tmp_path, monkeypatch):
    settings = configure_archive_test(monkeypatch, tmp_path)
    monkeypatch.setattr(settings, "max_archive_compression_ratio", 2.0)
    payload = make_zip({"bomb.txt": "A" * 20_000})

    with TestClient(app) as client:
        response = upload_zip(client, payload)

    assert response.status_code == 413
    assert response.json() == {
        "detail": "ZIP archive exceeds the configured compression-ratio limit"
    }


def test_zip_upload_counts_streamed_bytes_even_with_false_content_length(
    tmp_path,
    monkeypatch,
):
    settings = configure_archive_test(monkeypatch, tmp_path)
    monkeypatch.setattr(settings, "max_archive_upload_bytes", 64)
    payload = make_zip({"README.md": "x" * 100}, compression=zipfile.ZIP_STORED)

    with TestClient(app) as client:
        response = upload_zip(
            client,
            payload,
            headers={"Content-Length": "1"},
        )

    assert response.status_code == 413
    assert response.json() == {
        "detail": "ZIP upload exceeds the configured size limit"
    }


def test_zip_upload_requires_content_type_and_safe_decoded_name(
    tmp_path,
    monkeypatch,
):
    configure_archive_test(monkeypatch, tmp_path)
    payload = make_zip({"README.md": "# demo"})

    with TestClient(app) as client:
        wrong_type = client.post(
            "/repository/upload-zip",
            content=payload,
            headers={
                "Content-Type": "application/octet-stream",
                "X-Archive-Name": "demo.zip",
            },
        )
        unsafe_name = client.post(
            "/repository/upload-zip",
            content=payload,
            headers={
                "Content-Type": "application/zip",
                "X-Archive-Name": "folder%2Fdemo.zip",
            },
        )

    assert wrong_type.status_code == 415
    assert wrong_type.json() == {
        "detail": "Content-Type must be application/zip"
    }
    assert unsafe_name.status_code == 400
    assert unsafe_name.json() == {
        "detail": "X-Archive-Name must be a valid ZIP filename"
    }


def test_zip_upload_requires_csrf_for_browser_session(tmp_path, monkeypatch):
    settings = configure_archive_test(monkeypatch, tmp_path)
    monkeypatch.setattr(settings, "admin_username", "admin")
    monkeypatch.setattr(
        settings,
        "admin_password_hash",
        auth.hash_password("correct-password"),
    )
    monkeypatch.setattr(
        settings,
        "auth_session_secret",
        "test-session-secret-with-enough-entropy",
    )
    monkeypatch.setattr(settings, "auth_cookie_secure", False)
    payload = make_zip({"README.md": "# demo"})

    with TestClient(app) as client:
        login = client.post(
            "/auth/login",
            json={"username": "admin", "password": "correct-password"},
        )
        response = upload_zip(client, payload)

    assert login.status_code == 200
    assert response.status_code == 403
    assert response.json() == {"detail": "invalid or missing CSRF token"}


def test_zip_upload_reuses_global_import_slot(tmp_path, monkeypatch):
    configure_archive_test(monkeypatch, tmp_path)
    started = Event()
    release = Event()
    payload = make_zip({"README.md": "# demo"})

    def blocking_load(*_args, **_kwargs):
        started.set()
        assert release.wait(timeout=5)
        return "uploaded-demo-123", [], 1

    monkeypatch.setattr(main_module, "load_zip_repository", blocking_load)
    monkeypatch.setattr(main_module, "index_chunks_incremental", fake_index_result)

    with TestClient(app) as client, ThreadPoolExecutor(max_workers=1) as executor:
        first_future = executor.submit(upload_zip, client, payload)
        assert started.wait(timeout=5)
        try:
            second = upload_zip(client, payload, name="other.zip")
            assert second.status_code == 429
            assert second.json() == {
                "detail": "repository import capacity reached"
            }
        finally:
            release.set()

        assert first_future.result(timeout=5).status_code == 200
